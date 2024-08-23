import os
import streamlit as st
import pandas as pd
from docx import Document
import base64
import io
from openpyxl import load_workbook
from openpyxl.styles import Font
from google.cloud import translate_v3 as translate

DATASET_ID = "example-set"  # change this to any dataset ID

def create_adaptive_mt_dataset(source_lang, target_lang):
    client = translate.TranslationServiceClient()
    adaptive_mt_dataset = translate.types.AdaptiveMtDataset()
    adaptive_mt_dataset.name = f"projects/d-dcaf/locations/us-central1/adaptiveMtDatasets/{DATASET_ID}"
    adaptive_mt_dataset.display_name = "Example set"
    adaptive_mt_dataset.source_language_code = source_lang
    adaptive_mt_dataset.target_language_code = target_lang
    request = translate.CreateAdaptiveMtDatasetRequest(
        parent="projects/d-dcaf/locations/us-central1",
        adaptive_mt_dataset=adaptive_mt_dataset,
    )
    response = client.create_adaptive_mt_dataset(request=request)
    return response

def import_adaptive_mt_file(file_content, source_lang, target_lang):
    client = translate.TranslationServiceClient()
    # You need to implement the logic to prepare the request based on your file content
    # This is a placeholder and needs to be adjusted based on your specific requirements
    request = translate.ImportAdaptiveMtFileRequest(
        parent=f"projects/d-dcaf/locations/us-central1/adaptiveMtDatasets/{DATASET_ID}",
        source_language_code=source_lang,
        target_language_code=target_lang,
        # Add other necessary parameters
    )
    response = client.import_adaptive_mt_file(request)
    return response

def adaptive_mt_translate(text, source_lang, target_lang):
    client = translate.TranslationServiceClient()
    request = translate.AdaptiveMtTranslateRequest(
        parent="projects/d-dcaf/locations/us-central1",
        dataset=f"projects/d-dcaf/locations/us-central1/adaptiveMtDatasets/{DATASET_ID}",
        content=[text],
        source_language_code=source_lang,
        target_language_code=target_lang
    )
    response = client.adaptive_mt_translate(request)
    return response.translations[0].translated_text

def save_translated_file(translated_text, file_name):
    if not os.path.exists("done"):
        os.makedirs("done")
    file_path = os.path.join("done", file_name)
    if os.path.exists(file_path):
        file_name, file_extension = os.path.splitext(file_name)
        count = 1
        while os.path.exists(os.path.join("done", f"{file_name}_{count}{file_extension}")):
            count += 1
        file_path = os.path.join("done", f"{file_name}_{count}{file_extension}")
    
    with open(file_path, "w", encoding="utf-8") as file:
        file.write(translated_text)
    return file_path

def save_excel_file(translated_df, file_name):
    if not os.path.exists("done"):
        os.makedirs("done")
    file_path = os.path.join("done", file_name)
    if os.path.exists(file_path):
        file_name, file_extension = os.path.splitext(file_name)
        count = 1
        while os.path.exists(os.path.join("done", f"{file_name}_{count}{file_extension}")):
            count += 1
        file_path = os.path.join("done", f"{file_name}_{count}{file_extension}")
    
    translated_df.to_excel(file_path, index=False)
    return file_path

def save_docx_file(translated_doc, file_name):
    if not os.path.exists("done"):
        os.makedirs("done")
    file_path = os.path.join("done", file_name)
    if os.path.exists(file_path):
        file_name, file_extension = os.path.splitext(file_name)
        count = 1
        while os.path.exists(os.path.join("done", f"{file_name}_{count}{file_extension}")):
            count += 1
        file_path = os.path.join("done", f"{file_name}_{count}{file_extension}")
    
    translated_doc.save(file_path)
    return file_path

def create_download_link(file_path):
    with open(file_path, "rb") as file:
        b64 = base64.b64encode(file.read()).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{os.path.basename(file_path)}"><button style="background-color: #008CBA; border: none; color: white; padding: 10px 20px; text-align: center; text-decoration: none; display: inline-block; font-size: 16px; margin: 4px 2px; cursor: pointer; border-radius: 12px;">Download {os.path.basename(file_path)}</button></a>'
    return href

def translate_excel(file_path, source_lang, target_lang):
    df = pd.read_excel(file_path)
    num_rows, num_cols = df.shape
    translated_df = pd.DataFrame(columns=df.columns)

    for col in range(num_cols):
        translated_col = []
        for row in range(num_rows):
            cell_value = df.iloc[row, col]
            if pd.notnull(cell_value):
                translated_text = adaptive_mt_translate(str(cell_value), source_lang, target_lang)
                translated_col.append(translated_text)
            else:
                translated_col.append(None)
        translated_df[df.columns[col]] = translated_col

    return translated_df

def translate_docx_with_style(file_content, source_lang, target_lang):
    translated_doc = Document()
    doc = Document(io.BytesIO(file_content))

    for paragraph in doc.paragraphs:
        translated_paragraph = translated_doc.add_paragraph()
        for run in paragraph.runs:
            translated_text = adaptive_mt_translate(run.text, source_lang, target_lang)
            new_run = translated_paragraph.add_run(translated_text)
            
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.size = run.font.size
            new_run.font.name = run.font.name
            
        translated_paragraph.style = paragraph.style

    return translated_doc

def main():
    st.title("YOP2 TR")
    st.write("Upload File nya cuy")
    
    uploaded_file = st.file_uploader("hm", type=["txt", "xlsx", "docx"])
    target_language_options = ["Arabic", "German", "Spanish", "French", "Hindi", "Indonesian", "Japanese", "Chinese"]
    target_language_codes = ["ar", "de", "es", "fr", "hi", "id", "ja", "zh-CN"]

    source_language = st.selectbox("Source Language", ["English"] + target_language_options)
    target_language = st.selectbox("Diterjemahin kemana?", target_language_options)

    source_lang_code = "en" if source_language == "English" else target_language_codes[target_language_options.index(source_language)]
    target_lang_code = target_language_codes[target_language_options.index(target_language)]

    if 'translation_complete' not in st.session_state:
        st.session_state.translation_complete = False
    if 'translated_file_path' not in st.session_state:
        st.session_state.translated_file_path = None

    if uploaded_file is not None:
        if st.session_state.translation_complete:
            if st.button("Download"):
                if st.session_state.translated_file_path:
                    st.markdown(create_download_link(st.session_state.translated_file_path), unsafe_allow_html=True)
        else:
            if st.button("Mulai"):
                file_extension = uploaded_file.name.split(".")[-1]
                
                # Create adaptive MT dataset
                create_adaptive_mt_dataset(source_lang_code, target_lang_code)
                
                with st.spinner('Menerjemahkan...'):
                    if file_extension == "txt":
                        content = uploaded_file.read().decode()
                        translated_text = adaptive_mt_translate(content, source_lang_code, target_lang_code)
                        translated_file_path = save_translated_file(translated_text, f"translated_{uploaded_file.name}")
                    elif file_extension == "xlsx":
                        translated_df = translate_excel(uploaded_file, source_lang_code, target_lang_code)
                        translated_file_path = save_excel_file(translated_df, f"translated_{uploaded_file.name}")
                    elif file_extension == "docx":
                        translated_doc = translate_docx_with_style(uploaded_file.read(), source_lang_code, target_lang_code)
                        translated_file_path = save_docx_file(translated_doc, f"translated_{uploaded_file.name}")
                    else:
                        st.error("Format file tidak didukung. Silakan unggah file TXT, XLSX, atau DOCX.")
                        return

                st.success('Terjemahan selesai!')
                st.session_state.translation_complete = True
                st.session_state.translated_file_path = translated_file_path
                st.experimental_rerun()

if __name__ == "__main__":
    main()
